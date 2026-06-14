import streamlit as st
from PIL import Image, ImageStat, ImageEnhance, ImageOps
from io import BytesIO
import uuid, os, zipfile
import fitz
import pytesseract
from utils import inject_css, save_history

os.environ["U2NET_HOME"] = "/tmp"

st.set_page_config(page_title="SmartDoc Studio", page_icon="📄", layout="wide", initial_sidebar_state="expanded")
inject_css()

# ── Sidebar ───────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style="padding:0.5rem 0 1rem 0;">
        <div style="font-family:'Space Grotesk',sans-serif;font-size:1.1rem;font-weight:700;color:#1e293b;">📄 SmartDoc Studio</div>
    </div>
    """, unsafe_allow_html=True)
    st.page_link("app.py", label="← Home", icon="🏠")
    st.page_link("pages/2_🎬_VideoStudio.py", label="🎬 Video Studio", icon="🎬")
    st.divider()
    st.caption("SESSION HISTORY")
    if "history" not in st.session_state or not st.session_state.history:
        st.caption("No actions yet.")
    else:
        for item in reversed(st.session_state.history[-8:]):
            st.markdown(f'<div class="hist-badge">· <b>{item["name"][:22]}</b><br><span style="color:#94a3b8;font-size:0.78rem">{item["action"]}</span></div>', unsafe_allow_html=True)

# ── Config ────────────────────────────────────────────────────────
VISA_STANDARDS = {
    "US Visa (600×600)":    {"width": 600, "height": 600,  "min_brightness": 180},
    "UK Visa (600×750)":    {"width": 600, "height": 750,  "min_brightness": 180},
    "Schengen (350×450)":   {"width": 350, "height": 450,  "min_brightness": 175},
    "India (602×602)":      {"width": 602, "height": 602,  "min_brightness": 170},
    "UAE (300×369)":        {"width": 300, "height": 369,  "min_brightness": 175},
    "Canada (420×540)":     {"width": 420, "height": 540,  "min_brightness": 175},
}

IMAGE_FORMATS = {
    "JPEG (.jpg)":               ("JPEG", "jpg",  "image/jpeg"),
    "PNG (.png)":                ("PNG",  "png",  "image/png"),
    "WebP (.webp)":              ("WEBP", "webp", "image/webp"),
    "BMP (.bmp)":                ("BMP",  "bmp",  "image/bmp"),
    "TIFF (.tiff)":              ("TIFF", "tiff", "image/tiff"),
    "GIF (.gif)":                ("GIF",  "gif",  "image/gif"),
    "ICO (.ico)":                ("ICO",  "ico",  "image/x-icon"),
    "PPM (.ppm)":                ("PPM",  "ppm",  "image/x-portable-pixmap"),
    "TGA (.tga)":                ("TGA",  "tga",  "image/x-tga"),
    "PCX (.pcx)":                ("PCX",  "pcx",  "image/x-pcx"),
    "SGI (.sgi)":                ("SGI",  "sgi",  "image/sgi"),
    "Grayscale PNG":             ("PNG",  "png",  "image/png"),
    "Sepia JPEG":                ("JPEG", "jpg",  "image/jpeg"),
    "High-Res JPEG (300dpi)":   ("JPEG", "jpg",  "image/jpeg"),
    "Thumbnail 128px":           ("JPEG", "jpg",  "image/jpeg"),
    "Square Crop PNG":           ("PNG",  "png",  "image/png"),
    "Passport Size JPEG":        ("JPEG", "jpg",  "image/jpeg"),
    "Social Square 1080×1080":  ("PNG",  "png",  "image/png"),
    "Banner 1200×628":           ("PNG",  "png",  "image/png"),
    "Avatar 200×200":            ("PNG",  "png",  "image/png"),
}

ALLOWED = ["png", "jpg", "jpeg", "bmp", "tiff", "webp"]

BG_COLORS = {
    "White": (255,255,255), "Light Gray": (220,220,220), "Blue": (0,112,201),
    "Red": (220,50,50), "Green": (50,180,100), "Black": (0,0,0),
    "Yellow": (255,220,0), "Navy": (10,30,80), "Cream": (255,248,220),
}

# ── Helpers ───────────────────────────────────────────────────────
def pil_bytes(img, fmt="PNG"):
    buf = BytesIO()
    if fmt == "JPEG" and img.mode in ("RGBA","P"):
        img = img.convert("RGB")
    img.save(buf, format=fmt)
    return buf.getvalue()

def pdf_to_docx_visual(pdf_bytes):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Pt(18)
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    pw = sec.page_width.inches - 0.25
    for i, page in enumerate(pdf):
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        buf = BytesIO(pix.tobytes("png"))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(buf, width=Inches(pw))
        if i < len(pdf) - 1:
            doc.add_page_break()
    pdf.close()
    out = BytesIO(); doc.save(out); return out.getvalue()

def pdf_text_only(pdf_bytes):
    from docx import Document
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = "\n\n".join(f"--- Page {p.number+1} ---\n{p.get_text()}" for p in pdf)
    pdf.close()
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    out = BytesIO(); doc.save(out); return out.getvalue()

def convert_image(img, label, fmt):
    out = img.copy()
    if label == "Grayscale PNG":
        out = out.convert("L")
    elif label == "Sepia JPEG":
        g = out.convert("L")
        out = Image.merge("RGB",[g.point(lambda p:min(int(p*1.1),255)),g.point(lambda p:int(p*0.9)),g.point(lambda p:int(p*0.7))])
    elif label == "High-Res JPEG (300dpi)":
        buf=BytesIO(); out.convert("RGB").save(buf,format="JPEG",dpi=(300,300),quality=95); return buf.getvalue()
    elif label == "Thumbnail 128px":
        out.thumbnail((128,128),Image.LANCZOS)
    elif label == "Square Crop PNG":
        w,h=out.size; s=min(w,h); out=out.crop(((w-s)//2,(h-s)//2,(w-s)//2+s,(h-s)//2+s))
    elif label == "Passport Size JPEG":
        out=out.resize((413,531),Image.LANCZOS)
    elif label == "Social Square 1080×1080":
        out=ImageOps.fit(out,(1080,1080),Image.LANCZOS)
    elif label == "Banner 1200×628":
        out=ImageOps.fit(out,(1200,628),Image.LANCZOS)
    elif label == "Avatar 200×200":
        out=ImageOps.fit(out,(200,200),Image.LANCZOS)
    return pil_bytes(out, fmt)

def remove_bg(img):
    from rembg import remove
    out = remove(img)
    return out.convert("RGBA") if isinstance(out, Image.Image) else Image.open(BytesIO(out)).convert("RGBA")

def apply_bg(rgba, choice, custom=None):
    if choice == "Transparent": return rgba
    color = custom if choice == "Custom" else BG_COLORS.get(choice,(255,255,255))
    bg = Image.new("RGBA", rgba.size, color+(255,))
    bg.paste(rgba, mask=rgba.split()[3])
    return bg.convert("RGB")

def validate_visa(img, std):
    warns=[]
    w,h=img.size
    if (w,h)!=(std["width"],std["height"]): warns.append(f"Size must be {std['width']}×{std['height']} (yours: {w}×{h})")
    br=ImageStat.Stat(img.convert("L")).mean[0]
    if br<std["min_brightness"]: warns.append(f"Too dark (brightness {br:.0f}, need ≥{std['min_brightness']})")
    return len(warns)==0, warns

def fix_visa(img, std):
    img=img.convert("RGB")
    br=ImageStat.Stat(img.convert("L")).mean[0]
    if br<std["min_brightness"]:
        img=ImageEnhance.Brightness(img).enhance(min(std["min_brightness"]/max(br,1)*1.05,2.5))
    return ImageEnhance.Contrast(img).enhance(1.1).resize((std["width"],std["height"]),Image.LANCZOS)

def fix_visa_ai(img, std):
    rgba=remove_bg(img)
    bg=Image.new("RGBA",rgba.size,(255,255,255,255))
    bg.paste(rgba,mask=rgba.split()[3])
    return fix_visa(bg.convert("RGB"),std)

# ── UI ────────────────────────────────────────────────────────────
st.markdown('<div class="hero-title">📄 SmartDoc Studio</div>', unsafe_allow_html=True)
st.markdown('<div class="hero-sub">PDF → Word · Image Convert · Visa Photo · Background Removal</div>', unsafe_allow_html=True)

tabs = st.tabs(["📄 PDF → Word", "🖼️ Image Convert", "🪪 Visa Photo", "✂️ Background", "🕓 History"])

# ── PDF ───────────────────────────────────────────────────────────
with tabs[0]:
    st.markdown('<div class="card-title">PDF to Word — Layout-Preserving</div>', unsafe_allow_html=True)
    file = st.file_uploader("Upload PDF", type=["pdf"], key="pdf_up")
    if file:
        pdf_bytes = file.read()
        npages = len(fitz.open(stream=pdf_bytes, filetype="pdf"))
        st.info(f"📄 **{file.name}** — {npages} page(s)")
        mode = st.radio("Mode", ["🖼️ Visual (exact layout)", "🔤 Text only (fast)"], key="pdf_mode")
        if st.button("Convert to Word", key="pdf_btn"):
            with st.spinner("Converting..."):
                out = pdf_to_docx_visual(pdf_bytes) if "Visual" in mode else pdf_text_only(pdf_bytes)
                label = "visual" if "Visual" in mode else "text"
                save_history(file.name, f"PDF→DOCX ({label})")
            st.success("✅ Done!")
            st.download_button("⬇️ Download Word Document", out,
                f"{file.name.rsplit('.',1)[0]}_{label}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="pdf_dl")
        st.caption("💡 Visual mode preserves exact layout as images. Text mode extracts selectable text.")

# ── IMAGE CONVERT ────────────────────────────────────────────────
with tabs[1]:
    st.markdown('<div class="card-title">Image Converter — 20+ Formats & Presets</div>', unsafe_allow_html=True)
    file = st.file_uploader("Upload image", type=ALLOWED, key="img_up")
    if file:
        img = Image.open(file)
        w, h = img.size
        col1, col2 = st.columns([1,2])
        with col1:
            st.image(img, caption=f"{w}×{h}", use_container_width=True)
        with col2:
            selected = st.multiselect("Output formats", list(IMAGE_FORMATS.keys()), default=["PNG (.png)","JPEG (.jpg)"], key="img_fmts")
            st.slider("JPEG quality", 60, 100, 92, key="img_q")
        if selected and st.button("Convert", key="img_btn"):
            results = {}
            with st.spinner("Converting..."):
                for lbl in selected:
                    fmt, ext, mime = IMAGE_FORMATS[lbl]
                    results[lbl] = (convert_image(img, lbl, fmt), ext, mime)
                save_history(file.name, f"→ {', '.join(selected)}")
            st.success(f"✅ {len(results)} format(s) ready")
            if len(results) == 1:
                lbl,(data,ext,mime) = next(iter(results.items()))
                st.download_button(f"⬇️ Download {lbl}", data, f"converted.{ext}", mime, key="img_dl1")
            else:
                zbuf = BytesIO()
                with zipfile.ZipFile(zbuf,"w",zipfile.ZIP_DEFLATED) as zf:
                    for i,(lbl,(data,ext,mime)) in enumerate(results.items()):
                        zf.writestr(f"{i+1:02d}_{lbl.replace('/','').replace(' ','_')}.{ext}", data)
                st.download_button("⬇️ Download All as ZIP", zbuf.getvalue(), "converted.zip", "application/zip", key="img_zip")
                for i,(lbl,(data,ext,mime)) in enumerate(results.items()):
                    st.download_button(f"⬇️ {lbl}", data, f"converted_{i+1}.{ext}", mime, key=f"img_dl_{i}")

# ── VISA ──────────────────────────────────────────────────────────
with tabs[2]:
    st.markdown('<div class="card-title">Visa Photo Fixer</div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        std_name = st.selectbox("Visa standard", list(VISA_STANDARDS.keys()), key="visa_std")
        std = VISA_STANDARDS[std_name]
        st.caption(f"Size: **{std['width']}×{std['height']}px** · Min brightness: {std['min_brightness']}")
    with col2:
        file = st.file_uploader("Upload photo", type=ALLOWED, key="visa_up")
    if file:
        img = Image.open(file)
        c1, c2 = st.columns(2)
        with c1: st.image(img, caption="Original", use_container_width=True)
        method = st.radio("Fix method", ["🎨 Auto-fix (brightness + resize)", "🤖 AI Fix (remove BG + auto-fix)"], key="visa_method")
        cv, cf = st.columns(2)
        with cv:
            if st.button("✅ Validate", key="visa_val"):
                ok, warns = validate_visa(img, std)
                if ok: st.success("✅ Valid!")
                else:
                    for w in warns: st.warning(w)
        with cf:
            if st.button("🔧 Fix Photo", key="visa_fix"):
                with st.spinner("Fixing..."):
                    fixed = fix_visa_ai(img,std) if "AI" in method else fix_visa(img,std)
                    save_history(file.name, f"Visa Fix ({std_name})")
                with c2: st.image(fixed, caption="Fixed", use_container_width=True)
                buf=BytesIO(); fixed.save(buf,format="JPEG",quality=95,dpi=(300,300))
                st.download_button("⬇️ Download Visa Photo (JPEG 300dpi)", buf.getvalue(), "visa_photo.jpg", "image/jpeg", key="visa_dl")
                ok2,w2=validate_visa(fixed,std)
                if ok2: st.success("✅ Fixed photo meets all requirements")
                else:
                    for w in w2: st.warning(f"⚠️ {w}")

# ── BACKGROUND ────────────────────────────────────────────────────
with tabs[3]:
    st.markdown('<div class="card-title">Background Removal & Replacement</div>', unsafe_allow_html=True)
    file = st.file_uploader("Upload image", type=ALLOWED, key="bg_up")
    if file:
        img = Image.open(file)
        st.image(img, caption="Original", width=300)
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            all_bg = ["Transparent"] + list(BG_COLORS.keys()) + ["Custom"]
            bg_choice = st.selectbox("Background", all_bg, key="bg_choice")
        custom_rgb = None
        if bg_choice == "Custom":
            with c2:
                r = st.slider("R", 0, 255, 100, key="bg_r")
                g = st.slider("G", 0, 255, 180, key="bg_g")
                b = st.slider("B", 0, 255, 255, key="bg_b")
                custom_rgb = (r,g,b)
                st.color_picker("Color preview", f"#{r:02x}{g:02x}{b:02x}", key="bg_prev", disabled=True)
        if st.button("✂️ Remove Background", key="bg_btn"):
            with st.spinner("Running AI model..."):
                rgba = remove_bg(img)
                result = apply_bg(rgba, bg_choice, custom_rgb)
                save_history(file.name, f"BG → {bg_choice}")
            st.image(result, caption=f"Background: {bg_choice}", width=300)
            is_t = bg_choice == "Transparent"
            out_img = rgba if is_t else result
            buf=BytesIO(); out_img.save(buf,format="PNG")
            st.download_button("⬇️ Download PNG", buf.getvalue(), "no_bg.png", "image/png", key="bg_dl_png")
            if not is_t:
                bj=BytesIO(); result.convert("RGB").save(bj,format="JPEG",quality=95)
                st.download_button("⬇️ Download JPEG", bj.getvalue(), "no_bg.jpg", "image/jpeg", key="bg_dl_jpg")

# ── HISTORY ───────────────────────────────────────────────────────
with tabs[4]:
    st.markdown('<div class="card-title">Session History</div>', unsafe_allow_html=True)
    if "history" not in st.session_state or not st.session_state.history:
        st.info("No actions yet.")
    else:
        for item in reversed(st.session_state.history):
            st.markdown(f'<div class="hist-badge">📄 <b>{item["name"]}</b> → {item["action"]}</div>', unsafe_allow_html=True)
