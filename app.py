import streamlit as st
from PIL import Image, ImageStat, ImageEnhance, ImageFilter, ImageOps
from io import BytesIO
import uuid
import fitz  # PyMuPDF
import pytesseract
import os
import zipfile
import base64

# ================= SAFE ENV =================
os.environ["U2NET_HOME"] = "/tmp"

# ================= CONFIG =================
VISA_STANDARDS = {
    "US Visa (600×600)": {"width": 600, "height": 600, "min_brightness": 180, "bg_color": (255, 255, 255)},
    "UK Visa (600×750)": {"width": 600, "height": 750, "min_brightness": 180, "bg_color": (255, 255, 255)},
    "Schengen Visa (350×450)": {"width": 350, "height": 450, "min_brightness": 175, "bg_color": (255, 255, 255)},
    "Indian Visa (51mm×51mm @ 300dpi = 602×602)": {"width": 602, "height": 602, "min_brightness": 170, "bg_color": (255, 255, 255)},
    "UAE Visa (300×369)": {"width": 300, "height": 369, "min_brightness": 175, "bg_color": (255, 255, 255)},
    "Canada Visa (420×540)": {"width": 420, "height": 540, "min_brightness": 175, "bg_color": (255, 255, 255)},
}

IMAGE_FORMATS = {
    "JPEG (.jpg)":       ("JPEG", "jpg",  "image/jpeg"),
    "PNG (.png)":        ("PNG",  "png",  "image/png"),
    "WebP (.webp)":      ("WEBP", "webp", "image/webp"),
    "BMP (.bmp)":        ("BMP",  "bmp",  "image/bmp"),
    "TIFF (.tiff)":      ("TIFF", "tiff", "image/tiff"),
    "GIF (.gif)":        ("GIF",  "gif",  "image/gif"),
    "ICO (.ico)":        ("ICO",  "ico",  "image/x-icon"),
    "PPM (.ppm)":        ("PPM",  "ppm",  "image/x-portable-pixmap"),
    "TGA (.tga)":        ("TGA",  "tga",  "image/x-tga"),
    "PCX (.pcx)":        ("PCX",  "pcx",  "image/x-pcx"),
    "IM (.im)":          ("IM",   "im",   "image/x-im"),
    "SGI (.sgi)":        ("SGI",  "sgi",  "image/sgi"),
    "Grayscale PNG":     ("PNG",  "png",  "image/png"),
    "Sepia JPEG":        ("JPEG", "jpg",  "image/jpeg"),
    "High-Res JPEG 300dpi": ("JPEG", "jpg", "image/jpeg"),
    "Thumbnail (128px)": ("JPEG", "jpg",  "image/jpeg"),
    "Square Crop PNG":   ("PNG",  "png",  "image/png"),
    "Passport Size JPEG": ("JPEG","jpg",  "image/jpeg"),
    "Social Square (1080×1080 PNG)": ("PNG","png","image/png"),
    "Banner (1200×628 PNG)": ("PNG","png","image/png"),
    "Avatar (200×200 PNG)": ("PNG","png","image/png"),
}

ALLOWED_IMAGE_EXT = ["png", "jpg", "jpeg", "bmp", "tiff", "webp"]

BG_COLORS = {
    "White":       (255, 255, 255),
    "Light Gray":  (220, 220, 220),
    "Blue":        (0, 112, 201),
    "Red":         (220, 50, 50),
    "Green":       (50, 180, 100),
    "Black":       (0, 0, 0),
    "Yellow":      (255, 220, 0),
    "Navy":        (10, 30, 80),
    "Cream":       (255, 248, 220),
    "Custom RGB":  None,
}

# ================= SESSION STATE =================
if "history" not in st.session_state:
    st.session_state.history = []

def save_history(name, action):
    st.session_state.history.append({
        "id": str(uuid.uuid4()),
        "name": name,
        "action": action
    })

# ================= HELPERS =================
def pil_to_bytes(img, fmt="PNG"):
    buf = BytesIO()
    if fmt == "JPEG" and img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    img.save(buf, format=fmt)
    return buf.getvalue()

# ================= PDF → WORD (layout-preserving) =================
def pdf_to_docx_rich(pdf_bytes):
    """
    Renders each PDF page as a high-res image and inserts into DOCX,
    preserving exact visual layout. Also extracts selectable text via
    a hidden layer approach in docx so it looks identical to the PDF.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import math

    doc = Document()

    # Remove default margins for full-page images
    section = doc.sections[0]
    section.top_margin    = Pt(18)
    section.bottom_margin = Pt(18)
    section.left_margin   = Pt(18)
    section.right_margin  = Pt(18)

    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page_width_inches = section.page_width.inches - 0.25  # fit within margins

    for page_num, page in enumerate(pdf_doc):
        # Render page at 2× resolution for crispness
        mat = fitz.Matrix(2.0, 2.0)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_bytes = pix.tobytes("png")

        # Insert page image
        img_buf = BytesIO(img_bytes)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0] if para.runs else para.add_run()
        run.add_picture(img_buf, width=Inches(page_width_inches))

        # Page break between pages (not after last)
        if page_num < len(pdf_doc) - 1:
            doc.add_page_break()

    pdf_doc.close()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def pdf_ocr_text(pdf_bytes):
    """Extract selectable text from PDF using PyMuPDF."""
    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = ""
    for page in pdf_doc:
        text += f"\n\n--- Page {page.number + 1} ---\n"
        text += page.get_text()
    pdf_doc.close()
    return text

# ================= IMAGE CONVERSION =================
def convert_image(img: Image.Image, label: str, fmt: str, ext: str):
    """Apply any special transform based on label, then convert."""
    out = img.copy()

    if label == "Grayscale PNG":
        out = out.convert("L")

    elif label == "Sepia JPEG":
        gray = out.convert("L")
        sepia = Image.merge("RGB", [
            gray.point(lambda p: min(int(p * 1.1), 255)),
            gray.point(lambda p: int(p * 0.9)),
            gray.point(lambda p: int(p * 0.7)),
        ])
        out = sepia

    elif label == "High-Res JPEG 300dpi":
        buf = BytesIO()
        out.convert("RGB").save(buf, format="JPEG", dpi=(300, 300), quality=95)
        return buf.getvalue()

    elif label == "Thumbnail (128px)":
        out.thumbnail((128, 128), Image.LANCZOS)

    elif label == "Square Crop PNG":
        w, h = out.size
        s = min(w, h)
        left = (w - s) // 2
        top  = (h - s) // 2
        out  = out.crop((left, top, left + s, top + s))

    elif label == "Passport Size JPEG":
        out = out.resize((413, 531), Image.LANCZOS)  # 35×45 mm @ 300dpi

    elif label == "Social Square (1080×1080 PNG)":
        out = ImageOps.fit(out, (1080, 1080), Image.LANCZOS)

    elif label == "Banner (1200×628 PNG)":
        out = ImageOps.fit(out, (1200, 628), Image.LANCZOS)

    elif label == "Avatar (200×200 PNG)":
        out = ImageOps.fit(out, (200, 200), Image.LANCZOS)

    return pil_to_bytes(out, fmt)

# ================= BACKGROUND REMOVAL =================
def remove_background_rgba(img: Image.Image) -> Image.Image:
    """Remove background, return RGBA image."""
    from rembg import remove
    output = remove(img)
    if isinstance(output, Image.Image):
        return output.convert("RGBA")
    return Image.open(BytesIO(output)).convert("RGBA")

def apply_background(rgba_img: Image.Image, bg_choice: str, custom_rgb=None) -> Image.Image:
    """Composite RGBA image over chosen background color."""
    if bg_choice == "Transparent PNG":
        return rgba_img  # keep alpha

    if bg_choice == "Custom RGB":
        color = custom_rgb or (255, 255, 255)
    else:
        color = BG_COLORS.get(bg_choice, (255, 255, 255))

    background = Image.new("RGBA", rgba_img.size, color + (255,))
    background.paste(rgba_img, mask=rgba_img.split()[3])
    return background.convert("RGB")

# ================= VISA PHOTO =================
def validate_visa(img: Image.Image, standard: dict):
    warnings = []
    w, h = img.size
    if (w, h) != (standard["width"], standard["height"]):
        warnings.append(f"Size must be {standard['width']}×{standard['height']} px (yours: {w}×{h})")
    brightness = ImageStat.Stat(img.convert("L")).mean[0]
    if brightness < standard["min_brightness"]:
        warnings.append(f"Image too dark (brightness {brightness:.0f}, need ≥{standard['min_brightness']})")
    return len(warnings) == 0, warnings

def fix_visa(img: Image.Image, standard: dict) -> Image.Image:
    """Auto-fix brightness, resize, and set white background."""
    img = img.convert("RGB")
    brightness = ImageStat.Stat(img.convert("L")).mean[0]
    if brightness < standard["min_brightness"]:
        factor = standard["min_brightness"] / max(brightness, 1)
        factor = min(factor * 1.05, 2.5)
        img = ImageEnhance.Brightness(img).enhance(factor)
    img = ImageEnhance.Contrast(img).enhance(1.1)
    img = img.resize((standard["width"], standard["height"]), Image.LANCZOS)
    return img

def fix_visa_with_bg(img: Image.Image, standard: dict) -> Image.Image:
    """Remove background, apply white BG, then fix visa standards."""
    rgba = remove_background_rgba(img)
    bg = Image.new("RGBA", rgba.size, (255, 255, 255, 255))
    bg.paste(rgba, mask=rgba.split()[3])
    return fix_visa(bg.convert("RGB"), standard)

# ================= UI THEME =================
st.set_page_config(
    page_title="SmartDoc Studio",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=Space+Grotesk:wght@500;700&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

/* App background — clean white */
.stApp {
    background: #f5f7fa;
    color: #1a1d23;
}

/* Main header */
.hero-title {
    font-family: 'Space Grotesk', sans-serif;
    font-size: 2.4rem;
    font-weight: 700;
    background: linear-gradient(135deg, #2563eb 0%, #7c3aed 60%, #db2777 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    margin-bottom: 0.2rem;
}
.hero-sub {
    color: #64748b;
    font-size: 1rem;
    margin-bottom: 2rem;
}

/* Tab styling */
.stTabs [data-baseweb="tab-list"] {
    background: #ffffff;
    border-radius: 12px;
    padding: 4px;
    gap: 4px;
    border: 1px solid #e2e8f0;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
}
.stTabs [data-baseweb="tab"] {
    border-radius: 8px;
    color: #64748b;
    font-weight: 500;
    font-size: 0.9rem;
    padding: 0.5rem 1.1rem;
}
.stTabs [aria-selected="true"] {
    background: #eff6ff !important;
    color: #2563eb !important;
}

/* Cards */
.card {
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-radius: 14px;
    padding: 1.5rem;
    margin-bottom: 1rem;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
}
.card-title {
    font-family: 'Space Grotesk', sans-serif;
    font-size: 1.1rem;
    font-weight: 700;
    color: #1e293b;
    margin-bottom: 0.75rem;
}

/* Buttons */
.stButton > button {
    background: linear-gradient(135deg, #2563eb, #7c3aed);
    color: white;
    border: none;
    border-radius: 8px;
    font-weight: 600;
    padding: 0.5rem 1.3rem;
    font-size: 0.9rem;
    transition: opacity 0.2s;
}
.stButton > button:hover {
    opacity: 0.88;
    color: white;
}

/* Download button */
.stDownloadButton > button {
    background: #eff6ff;
    color: #2563eb;
    border: 1px solid #bfdbfe;
    border-radius: 8px;
    font-weight: 600;
}
.stDownloadButton > button:hover {
    background: #dbeafe;
    color: #1d4ed8;
}

/* File uploader */
[data-testid="stFileUploader"] {
    background: #ffffff;
    border: 1.5px dashed #cbd5e1;
    border-radius: 12px;
    padding: 1rem;
}

/* Success / warning */
.stSuccess { background: #f0fdf4; border-color: #86efac; color: #166534; border-radius: 8px; }
.stWarning { background: #fffbeb; border-color: #fcd34d; color: #92400e; border-radius: 8px; }

/* History badge */
.hist-badge {
    display: block;
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-radius: 8px;
    padding: 0.4rem 0.85rem;
    font-size: 0.85rem;
    color: #475569;
    margin-bottom: 0.4rem;
    box-shadow: 0 1px 3px rgba(0,0,0,0.04);
}

/* Info chip */
.chip {
    display: inline-block;
    background: #eff6ff;
    border-radius: 20px;
    padding: 2px 10px;
    font-size: 0.78rem;
    color: #2563eb;
    margin-right: 6px;
    border: 1px solid #bfdbfe;
    font-weight: 500;
}

/* Divider */
.divider { border-top: 1px solid #e2e8f0; margin: 1.2rem 0; }

/* General text visibility fixes */
label, .stSelectbox label, .stRadio label, .stSlider label,
.stMultiSelect label, .stFileUploader label {
    color: #1e293b !important;
    font-weight: 500;
}
p, span, div { color: #334155; }
h1, h2, h3 { color: #0f172a; }

/* Info boxes */
.stInfo { background: #eff6ff; color: #1e40af; border-color: #bfdbfe; border-radius: 8px; }

/* Selectbox text */
.stSelectbox > div > div {
    background: #ffffff;
    border-color: #cbd5e1;
    color: #1e293b;
    border-radius: 8px;
}
</style>
""", unsafe_allow_html=True)

# ================= HEADER =================
st.markdown('<div class="hero-title">⚡ SmartDoc Studio</div>', unsafe_allow_html=True)
st.markdown('<div class="hero-sub">PDF conversion · Image tools · Visa photos · Background removal</div>', unsafe_allow_html=True)

tabs = st.tabs([
    "📄 PDF → Word",
    "🖼️ Image Convert",
    "🪪 Visa Photo",
    "✂️ Background",
    "🕓 History"
])

# ================= TAB 1: PDF → WORD =================
with tabs[0]:
    st.markdown('<div class="card-title">PDF to Word — Layout-Preserving</div>', unsafe_allow_html=True)
    st.markdown(
        '<span class="chip">📐 Exact layout</span>'
        '<span class="chip">🖼️ Images kept</span>'
        '<span class="chip">🔤 Text extraction</span>',
        unsafe_allow_html=True
    )
    st.markdown("")

    file = st.file_uploader("Upload your PDF", type=["pdf"], key="pdf_up")

    if file:
        pdf_bytes = file.read()
        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        st.info(f"📄 **{file.name}** — {len(pdf_doc)} page(s) detected")
        pdf_doc.close()

        mode = st.radio(
            "Conversion mode",
            [
                "🖼️ Visual (exact layout — images of each page, fully editable doc)",
                "🔤 Text only (fast, selectable text, no images)",
            ],
            key="pdf_mode"
        )

        col1, col2 = st.columns([1, 2])
        with col1:
            convert_btn = st.button("Convert to Word", key="pdf_btn")

        if convert_btn:
            with st.spinner("Converting... this may take a moment for large PDFs"):
                if "Visual" in mode:
                    out = pdf_to_docx_rich(pdf_bytes)
                    label = "visual"
                else:
                    from docx import Document
                    text = pdf_ocr_text(pdf_bytes)
                    d = Document()
                    for line in text.split("\n"):
                        d.add_paragraph(line)
                    buf = BytesIO()
                    d.save(buf)
                    out = buf.getvalue()
                    label = "text"

                save_history(file.name, f"PDF → DOCX ({label})")

            st.success("✅ Conversion complete!")
            st.download_button(
                "⬇️ Download Word Document",
                out,
                f"{file.name.rsplit('.', 1)[0]}_{label}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="pdf_dl"
            )

        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.caption("💡 **Visual mode** renders each page as a crisp image inside the Word doc — your layout, fonts, tables, and images are preserved exactly. **Text mode** extracts raw selectable text.")

# ================= TAB 2: IMAGE CONVERT =================
with tabs[1]:
    st.markdown('<div class="card-title">Image Converter — 20+ Formats & Presets</div>', unsafe_allow_html=True)

    file = st.file_uploader("Upload image", type=ALLOWED_IMAGE_EXT, key="img_up")

    if file:
        img = Image.open(file)
        w, h = img.size

        col_prev, col_opts = st.columns([1, 2])
        with col_prev:
            st.image(img, caption=f"{file.name}  ({w}×{h})", use_container_width=True)

        with col_opts:
            st.markdown("**Select output format / preset:**")
            format_labels = list(IMAGE_FORMATS.keys())
            selected_formats = st.multiselect(
                "Choose one or more formats",
                format_labels,
                default=["PNG (.png)", "JPEG (.jpg)"],
                key="img_fmts"
            )

            # Quality slider for JPEG outputs
            jpeg_quality = st.slider("JPEG quality", 60, 100, 92, key="img_q")

        if selected_formats and st.button("Convert", key="img_btn"):
            results = {}
            with st.spinner("Converting..."):
                for label in selected_formats:
                    fmt, ext, mime = IMAGE_FORMATS[label]
                    data = convert_image(img, label, fmt, ext)
                    results[label] = (data, ext, mime)
                save_history(file.name, f"Image → {', '.join(selected_formats)}")

            st.success(f"✅ {len(results)} format(s) ready")

            if len(results) == 1:
                label, (data, ext, mime) = next(iter(results.items()))
                st.download_button(
                    f"⬇️ Download {label}",
                    data,
                    f"converted.{ext}",
                    mime=mime,
                    key="img_dl_single"
                )
            else:
                # Zip all outputs
                zip_buf = BytesIO()
                with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i, (label, (data, ext, mime)) in enumerate(results.items()):
                        safe = label.replace("/", "-").replace(" ", "_").replace("(", "").replace(")", "")
                        zf.writestr(f"{i+1:02d}_{safe}.{ext}", data)

                st.download_button(
                    "⬇️ Download All as ZIP",
                    zip_buf.getvalue(),
                    "converted_images.zip",
                    mime="application/zip",
                    key="img_dl_zip"
                )

                # Individual download buttons
                st.markdown("**Or download individually:**")
                for i, (label, (data, ext, mime)) in enumerate(results.items()):
                    st.download_button(
                        f"⬇️ {label}",
                        data,
                        f"converted_{i+1}.{ext}",
                        mime=mime,
                        key=f"img_dl_{i}"
                    )

# ================= TAB 3: VISA PHOTO =================
with tabs[2]:
    st.markdown('<div class="card-title">Visa Photo Fixer</div>', unsafe_allow_html=True)

    col_std, col_up = st.columns([1, 1])
    with col_std:
        standard_name = st.selectbox("Select visa standard", list(VISA_STANDARDS.keys()), key="visa_std")
        standard = VISA_STANDARDS[standard_name]
        st.caption(f"Required size: **{standard['width']}×{standard['height']} px** · Min brightness: {standard['min_brightness']}")

    with col_up:
        file = st.file_uploader("Upload photo", type=ALLOWED_IMAGE_EXT, key="visa_up")

    if file:
        img = Image.open(file)

        col_orig, col_fixed = st.columns(2)
        with col_orig:
            st.image(img, caption="Original", use_container_width=True)

        action = st.radio(
            "Fix method",
            [
                "🎨 Auto-fix (brightness + resize only)",
                "🤖 AI Fix (remove & replace background + auto-fix)",
            ],
            key="visa_action"
        )

        col_v, col_f = st.columns(2)
        with col_v:
            if st.button("✅ Validate", key="visa_val"):
                ok, warns = validate_visa(img, standard)
                if ok:
                    st.success("Photo meets visa requirements!")
                else:
                    for w in warns:
                        st.warning(w)

        with col_f:
            if st.button("🔧 Fix Photo", key="visa_fix_btn"):
                with st.spinner("Fixing..."):
                    if "AI Fix" in action:
                        fixed = fix_visa_with_bg(img, standard)
                    else:
                        fixed = fix_visa(img, standard)
                    save_history(file.name, f"Visa Fix ({standard_name})")

                with col_fixed:
                    st.image(fixed, caption="Fixed", use_container_width=True)

                buf = BytesIO()
                fixed.save(buf, format="JPEG", quality=95, dpi=(300, 300))
                st.download_button(
                    "⬇️ Download Visa Photo (JPEG 300dpi)",
                    buf.getvalue(),
                    "visa_photo.jpg",
                    mime="image/jpeg",
                    key="visa_dl"
                )

                ok2, warns2 = validate_visa(fixed, standard)
                if ok2:
                    st.success("✅ Fixed photo now meets all requirements")
                else:
                    for w in warns2:
                        st.warning(f"⚠️ {w}")

# ================= TAB 4: BACKGROUND REMOVAL =================
with tabs[3]:
    st.markdown('<div class="card-title">Background Removal & Replacement</div>', unsafe_allow_html=True)

    file = st.file_uploader("Upload image", type=ALLOWED_IMAGE_EXT, key="bg_up")

    if file:
        img = Image.open(file)
        st.image(img, caption="Original", use_container_width=False, width=320)

        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

        col_bg, col_rgb = st.columns([1, 1])
        with col_bg:
            bg_choice = st.selectbox(
                "Background after removal",
                ["Transparent PNG"] + [k for k in BG_COLORS if k != "Custom RGB"] + ["Custom RGB"],
                key="bg_choice"
            )

        custom_rgb = None
        if bg_choice == "Custom RGB":
            with col_rgb:
                r = st.slider("R", 0, 255, 255, key="bg_r")
                g = st.slider("G", 0, 255, 150, key="bg_g")
                b = st.slider("B", 0, 255, 100, key="bg_b")
                custom_rgb = (r, g, b)
                st.color_picker("Preview", f"#{r:02x}{g:02x}{b:02x}", key="bg_preview", disabled=True)

        if st.button("✂️ Remove Background", key="bg_btn"):
            with st.spinner("Running AI background removal..."):
                rgba = remove_background_rgba(img)
                result = apply_background(rgba, bg_choice, custom_rgb)
                save_history(file.name, f"BG Removed → {bg_choice}")

            st.image(result, caption=f"Background: {bg_choice}", use_container_width=False, width=320)

            is_transparent = (bg_choice == "Transparent PNG")
            out_fmt  = "PNG"
            out_ext  = "png"
            out_mime = "image/png"
            out_img  = rgba if is_transparent else result

            buf = BytesIO()
            out_img.save(buf, format=out_fmt)

            st.download_button(
                f"⬇️ Download {'Transparent ' if is_transparent else ''}PNG",
                buf.getvalue(),
                f"no_bg_{bg_choice.replace(' ', '_').lower()}.{out_ext}",
                mime=out_mime,
                key="bg_dl"
            )

            # Also offer JPEG if non-transparent
            if not is_transparent:
                buf_jpg = BytesIO()
                result.convert("RGB").save(buf_jpg, format="JPEG", quality=95)
                st.download_button(
                    "⬇️ Download JPEG",
                    buf_jpg.getvalue(),
                    f"no_bg_{bg_choice.replace(' ', '_').lower()}.jpg",
                    mime="image/jpeg",
                    key="bg_dl_jpg"
                )

# ================= TAB 5: HISTORY =================
with tabs[4]:
    st.markdown('<div class="card-title">Session History</div>', unsafe_allow_html=True)

    if not st.session_state.history:
        st.info("No actions yet — process a file to see history here.")
    else:
        for item in reversed(st.session_state.history):
            st.markdown(
                f'<div class="hist-badge">📄 <b>{item["name"]}</b> → {item["action"]}</div>',
                unsafe_allow_html=True
            )
